import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert


def replace_text_variables(participant_date):
    template = "report_template/report_template.docx"   
    doc = Document(template)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for key, value in participant_date.items():
                if key in run.text:
                    run.text = run.text.replace(key, value)
                    
    return doc


def add_images_to_end_of_document(doc):
    images = [
    'report_template/pupil_movement_plot.png',
    'report_template/threshold.png',
    'report_template/tracking.png'
    ]

    for image_path in images: # Add an images to the end of the document
        # Add an image to the end of the document
        p = doc.add_paragraph()
        r = p.add_run()
        r.add_picture(image_path, width=Inches(6))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    remove_images(images)


def remove_images(images):
    for image in images:
        os.remove(image)


def create_output_report(participant_data):
    directory = "reports"
    doc = replace_text_variables(participant_data)
    add_images_to_end_of_document(doc)

    if not os.path.exists(directory):
        os.makedirs(directory)
    
    doc.save(f"reports/{participant_data['{{name}}']}.docx")
    convert(f"reports/{participant_data['{{name}}']}.docx")

    